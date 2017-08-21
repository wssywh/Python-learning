#coding:utf-8
from docx import Document
from bs4 import BeautifulSoup
import urllib.request
import requests


url = "http://freebuf.com/news/94263.html"
# data = urllib.urlopen(url)
data = requests.get(url).content
document = Document()

soup = BeautifulSoup(data)
article = soup.find(name="div", attrs={'id': 'contenttxt'}).children
for e in article:
    try:
        if e.img:
            pic_name = ''
            print(e.img.attrs['src'])
            print(e.img)
            if 'gif' in e.img.attrs['src']:
                pic_name = 'temp.gif'
            elif 'png' in e.img.attrs['src']:
                pic_name = 'temp.png'
            elif 'jpg' in e.img.attrs['src']:
                pic_name = 'temp.jpg'
            else:
                pic_name = 'temp.jpeg'
            urllib.request.urlretrieve(e.img.attrs['src'], filename=pic_name)
            document.add_picture(pic_name)
    except:
        pass
    if e.string:
        print(e.string.encode('gbk', 'ignore'))

        document.add_paragraph(e.string)

document.save("freebuf_article.docx")
print("success create a document")
